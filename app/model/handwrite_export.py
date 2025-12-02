"""
Експорт розпізнаного тексту у різні формати
handwrite2print/app/model/export.py
"""
import os
from datetime import datetime


class TextExporter:
    """Клас для експорту тексту у різні формати"""
    
    def __init__(self):
        """Ініціалізація експортера.
        
        Метод порожній, оскільки клас не потребує початкової конфігурації.
        Всі необхідні налаштування виконуються в методах експорту.
        """
        pass
        
    def export_txt(self, text, file_path):
        """
        Експорт у текстовий файл
        
        Args:
            text: текст для експорту
            file_path: шлях до файлу
            
        Returns:
            True при успіху, False при помилці
        """
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
        except Exception as e:
            print(f"Помилка при збереженні TXT: {e}")
            return False
            
    def export_docx(self, text, file_path):
        """
        Експорт у Word документ
        
        Args:
            text: текст для експорту
            file_path: шлях до файлу
            
        Returns:
            True при успіху, False при помилці
        """
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Додавання заголовку
            doc.add_heading('Розпізнаний текст', level=1)
            
            # Додавання дати
            date_paragraph = doc.add_paragraph()
            date_paragraph.add_run(
                f'Дата створення: {datetime.now().strftime("%d.%m.%Y %H:%M")}'
            ).italic = True
            
            # Додавання тексту
            for line in text.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    # Налаштування шрифту
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                else:
                    doc.add_paragraph()  # Порожній рядок
                    
            # Збереження
            doc.save(file_path)
            return True
            
        except ImportError:
            print("Модуль python-docx не встановлено. Встановіть: pip install python-docx")
            return False
        except Exception as e:
            print(f"Помилка при збереженні DOCX: {e}")
            return False
            
    def export_pdf(self, text, file_path):
        """
        Експорт у PDF файл з підтримкою української кирилиці
        
        Args:
            text: текст для експорту
            file_path: шлях до файлу
            
        Returns:
            True при успіху, False при помилці
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            from reportlab.lib.colors import HexColor
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import platform
            
            # Створення PDF
            doc = SimpleDocTemplate(
                file_path,
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            story = []
            styles = getSampleStyleSheet()
            
            # Реєстрація шрифту, який підтримує кирилицю
            font_name = 'CyrillicFont'  # За замовчуванням
            font_registered = False
            
            try:
                system = platform.system()

                if system == 'Windows':
                    # Windows має шрифти, які підтримують кирилицю
                    # Шукаємо шрифт у стандартних місцях Windows
                    font_paths = [
                        'C:/Windows/Fonts/arial.ttf',
                        'C:/Windows/Fonts/ARIAL.TTF',  # Великі літери
                        'C:/Windows/Fonts/arialbd.ttf',  # Arial Bold
                        'C:/Windows/Fonts/times.ttf',
                        'C:/Windows/Fonts/timesnr.ttf',
                        'C:/Windows/Fonts/TIMES.TTF',  # Великі літери
                        'C:/Windows/Fonts/timesbd.ttf',  # Times New Roman Bold
                        'C:/Windows/Fonts/calibri.ttf',
                        'C:/Windows/Fonts/CALIBRI.TTF',  # Великі літери
                        'C:/Windows/Fonts/calibrib.ttf',  # Calibri Bold
                        'C:/Windows/Fonts/segoeui.ttf',  # Segoe UI
                        'C:/Windows/Fonts/SEGOEUI.TTF',  # Segoe UI великі літери
                    ]
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            try:
                                # Реєструємо шрифт з явним вказанням кодування
                                # Використовуємо subfontIndex=0 для правильного вибору підшрифту
                                font = TTFont('CyrillicFont', font_path, subfontIndex=0)
                                pdfmetrics.registerFont(font)
                                # Перевіряємо, чи шрифт дійсно зареєстровано
                                registered_fonts = pdfmetrics.getRegisteredFontNames()
                                if 'CyrillicFont' in registered_fonts:
                                    font_registered = True
                                    print(f"Зареєстровано шрифт для кирилиці: {font_path}")
                                    break
                                else:
                                    print(f"Шрифт {font_path} не з'явився в списку зареєстрованих")
                            except Exception as font_error:
                                print(f"Не вдалося зареєструвати шрифт {font_path}: {font_error}")
                                import traceback
                                print(traceback.format_exc())
                                continue
                
                elif system == 'Linux':
                    # Linux - спробуємо DejaVu Sans
                    font_paths = [
                        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
                        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                        '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
                    ]
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            try:
                                pdfmetrics.registerFont(TTFont('CyrillicFont', font_path))
                                font_registered = True
                                print(f"Зареєстровано шрифт для кирилиці: {font_path}")
                                break
                            except Exception as font_error:
                                print(f"Не вдалося зареєструвати шрифт {font_path}: {font_error}")
                                continue
                
                elif system == 'Darwin':  # macOS
                    font_paths = [
                        '/System/Library/Fonts/Supplemental/Arial.ttf',
                        '/System/Library/Fonts/Supplemental/Arial Bold.ttf',
                        '/Library/Fonts/Arial.ttf',
                        '/System/Library/Fonts/Helvetica.ttc',
                    ]
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            try:
                                pdfmetrics.registerFont(TTFont('CyrillicFont', font_path))
                                font_registered = True
                                print(f"Зареєстровано шрифт для кирилиці: {font_path}")
                                break
                            except Exception as font_error:
                                print(f"Не вдалося зареєструвати шрифт {font_path}: {font_error}")
                                continue
                
                # Якщо не вдалося зареєструвати шрифт, використовуємо fallback
                if not font_registered:
                    print("Попередження: не вдалося зареєструвати шрифт для кирилиці, використовується fallback")
                    # Спробуємо використати CID шрифти, які краще підтримують Unicode
                    try:
                        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                        # Спробуємо різні CID шрифти, які підтримують кирилицю
                        cid_fonts = ['HeiseiMin-W3', 'HeiseiKakuGo-W5', 'STSong-Light']
                        for cid_font in cid_fonts:
                            try:
                                pdfmetrics.registerFont(UnicodeCIDFont(cid_font))
                                font_name = cid_font
                                font_registered = True
                                print(f"Зареєстровано CID шрифт: {cid_font}")
                                break
                            except Exception:
                                continue
                    except Exception as cid_error:
                        print(f"Не вдалося зареєструвати CID шрифт: {cid_error}")
                    
                    # Якщо CID шрифти не спрацювали, використовуємо Helvetica з попередженням
                    if not font_registered:
                        font_name = 'Helvetica'
                        print("Попередження: використовується Helvetica, який може не підтримувати кирилицю")
                        
            except Exception as e:
                print(f"Помилка при реєстрації шрифту для кирилиці: {e}")
                font_name = 'Helvetica'
                font_registered = False
                
            # Стиль заголовку
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=HexColor("#2196F3"),
                spaceAfter=12,
                alignment=TA_CENTER,
                fontName=font_name
            )
            
            # Стиль дати
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=HexColor("#000000"),
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName=font_name
            )
            
            # Стиль основного тексту
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=12,
                leading=16,
                alignment=TA_LEFT,
                fontName=font_name
            )
            
            # Додавання заголовку
            story.append(Paragraph("Розпізнаний текст", title_style))
            
            # Додавання дати
            date_text = f"Дата створення: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            story.append(Paragraph(date_text, date_style))
            
            story.append(Spacer(1, 0.5*cm))
            
            # Додавання тексту
            for line in text.split('\n'):
                if line.strip():
                    # Переконуємося, що текст правильно закодований
                    # Конвертуємо в рядок, якщо потрібно
                    if not isinstance(line, str):
                        line = str(line)
                    
                    # Екранування спеціальних символів для ReportLab
                    # ReportLab вимагає екранування XML-символів
                    safe_line = (
                        line.replace('&', '&amp;')
                        .replace('<', '&lt;')
                        .replace('>', '&gt;')
                        .replace('"', '&quot;')
                        .replace("'", '&#39;')
                    )
                    
                    # Додаємо текст до PDF
                    # Використовуємо Paragraph з зареєстрованим шрифтом
                    try:
                        # Завжди використовуємо Paragraph з правильним шрифтом
                        if font_registered:
                            story.append(Paragraph(safe_line, body_style))
                        else:
                            # Якщо шрифт не зареєстровано, використовуємо Preformatted
                            # який краще обробляє Unicode символи навіть без спеціального шрифту
                            from reportlab.platypus import Preformatted
                            preformatted_style = ParagraphStyle(
                                'PreformattedStyle',
                                parent=body_style,
                                fontName='Courier'  # Courier зазвичай краще підтримує Unicode
                            )
                            story.append(Preformatted(line, preformatted_style, maxLineLength=80))
                        story.append(Spacer(1, 0.2 * cm))
                    except Exception as e:
                        # Якщо не вдалося додати параграф, спробуємо Preformatted
                        print(f"Помилка при додаванні рядка в PDF через Paragraph: {e}")
                        try:
                            from reportlab.platypus import Preformatted
                            preformatted_style = ParagraphStyle(
                                'PreformattedStyle',
                                parent=body_style,
                                fontName='Courier'  # Courier зазвичай краще підтримує Unicode
                            )
                            story.append(Preformatted(line, preformatted_style, maxLineLength=80))
                            story.append(Spacer(1, 0.2 * cm))
                        except Exception as preformatted_error:
                            print(f"Помилка при додаванні рядка через Preformatted: {preformatted_error}")
                            print(f"Не вдалося додати рядок: {line[:50]}...")
                else:
                    story.append(Spacer(1, 0.3 * cm))
                    
            # Побудова PDF
            doc.build(story)
            return True
            
        except ImportError:
            print("Модуль reportlab не встановлено. Встановіть: pip install reportlab")
            return False
        except Exception as e:
            print(f"Помилка при збереженні PDF: {e}")
            return False
            
    def export_html(self, text, file_path):
        """
        Експорт у HTML файл
        
        Args:
            text: текст для експорту
            file_path: шлях до файлу
            
        Returns:
            True при успіху, False при помилці
        """
        try:
            html_content = f"""
<!DOCTYPE html>
<html lang="uk">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Розпізнаний текст</title>
    <style>
        body {{
            font-family: 'Times New Roman', Times, serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        h1 {{
            color: #2196F3;
            text-align: center;
        }}
        .date {{
            text-align: center;
            color: #666;
            font-style: italic;
            margin-bottom: 30px;
        }}
        .content {{
            text-align: justify;
            white-space: pre-wrap;
        }}
    </style>
</head>
<body>
    <h1>Розпізнаний текст</h1>
    <p class="date">Дата створення: {datetime.now().strftime('%d.%m.%Y %H:%M')}</p>
    <div class="content">{self._escape_html(text)}</div>
</body>
</html>
"""
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            return True
            
        except Exception as e:
            print(f"Помилка при збереженні HTML: {e}")
            return False
            
    def _escape_html(self, text):
        """Екранування HTML спецсимволів"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))
